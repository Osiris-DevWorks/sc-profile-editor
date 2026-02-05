"""
Main application window
"""

from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                              QPushButton, QLabel, QFileDialog, QMessageBox,
                              QTextEdit, QTableWidget, QTableWidgetItem, QSplitter,
                              QLineEdit, QComboBox, QGroupBox, QCheckBox, QTabWidget,
                              QStyledItemDelegate, QDialog, QTextBrowser, QInputDialog,
                              QListWidget, QListWidgetItem, QSystemTrayIcon, QMenu)
from PyQt6.QtCore import Qt, QSortFilterProxyModel, QTimer, QUrl, QEvent
from PyQt6.QtGui import QStandardItemModel, QStandardItem, QDesktopServices, QPixmap, QCursor, QIcon
from PyQt6.QtGui import QAction
import sys
import os
import logging

logger = logging.getLogger(__name__)


def get_resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        # If not running as PyInstaller bundle, use the project root
        base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    return os.path.join(base_path, relative_path)

# Add parent directory to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from parser.xml_parser import ProfileParser
from parser.label_generator import LabelGenerator
from models.profile_model import ControlProfile
from utils.settings import AppSettings
from utils.version import get_version
from utils.device_splitter import get_device_for_input

# Import PDF widget - QtPdf is the only PDF viewer (WebEngine removed to reduce installer size)
PDF_WIDGET_AVAILABLE = False
PDF_WIDGET_CLASS = None

try:
    from gui.qtpdf_device_widget import QtPdfDeviceWidget
    PDF_WIDGET_CLASS = QtPdfDeviceWidget
    PDF_WIDGET_AVAILABLE = True
    logger.info("Using QtPdf for device view (native Qt PDF viewer)")
except ImportError as e:
    PDF_WIDGET_AVAILABLE = False
    logger.error(f"QtPdf widget not available: {e}")
    logger.error("Device View will be disabled")


class SelectAllDelegate(QStyledItemDelegate):
    """Custom delegate - clean editing with proper text replacement"""

    def __init__(self, parent, main_window):
        super().__init__(parent)
        self.main_window = main_window

    def createEditor(self, parent, option, index):
        """Create editor with proper configuration"""
        editor = QLineEdit(parent)
        # Force solid white background and black text
        editor.setStyleSheet("QLineEdit { background-color: white; color: black; }")
        return editor

    def setEditorData(self, editor, index):
        """Set editor data and select all for easy replacement"""
        if isinstance(editor, QLineEdit):
            # Get current text
            text = index.model().data(index, Qt.ItemDataRole.DisplayRole)
            editor.setText(text or "")
            # Ensure editor has focus before selecting
            editor.setFocus()
            # Use a small delay to ensure editor is ready
            QTimer.singleShot(0, editor.selectAll)
        else:
            super().setEditorData(editor, index)


class MainWindow(QMainWindow):
    """Main application window for SC Profile Editor"""

    def __init__(self):
        super().__init__()
        self.version = get_version()
        self.setWindowTitle(f"Star Citizen Profile Editor v{self.version}")
        self.setGeometry(100, 100, 1200, 800)

        # Initialize settings manager
        self.settings = AppSettings()

        # Store current profile
        self.current_profile = None
        self.current_profile_path = None

        # Store all bindings for filtering
        self.all_bindings = []

        # Input detection for key filter
        self.input_detector = None  # Lazy initialized when first used
        self.detecting_input_for_filter = False  # Track detection state

        # Initialize system tray icon
        self.tray_icon = None
        self.is_closing = False  # Flag to distinguish close from minimize
        self.setup_tray_icon()

        # Set window icon
        icon_path = get_resource_path("assets/icon.ico")
        self.setWindowIcon(QIcon(icon_path))

        # Create UI
        self.setup_ui()

        # Restore window geometry if saved
        self.restore_window_state()

        # Auto-load last profile after UI is shown (use QTimer to defer until after show())
        QTimer.singleShot(100, self.auto_load_last_profile)

    def setup_ui(self):
        """Set up the user interface"""
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Main layout
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)

        # Header
        header_layout = QHBoxLayout()
        title_label = QLabel(f"Star Citizen Profile Editor v{self.version}")
        title_label.setStyleSheet("font-size: 24px; font-weight: bold; margin: 10px;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()

        # Import button (leftmost)
        import_btn = QPushButton("Import Profile XML")
        import_btn.setStyleSheet("QPushButton { padding: 10px 20px; font-size: 14px; background-color: #4CAF50; color: white; border: none; border-radius: 4px; } QPushButton:hover { background-color: #45a049; }")
        import_btn.clicked.connect(self.import_profile)
        header_layout.addWidget(import_btn)

        # New Profile button
        new_profile_btn = QPushButton("New Profile")
        new_profile_btn.setStyleSheet("QPushButton { padding: 10px 20px; font-size: 14px; background-color: #2196F3; color: white; border: none; border-radius: 4px; } QPushButton:hover { background-color: #0b7dda; }")
        new_profile_btn.clicked.connect(self.create_new_profile)
        header_layout.addWidget(new_profile_btn)

        # Load Preset Profile button
        load_preset_btn = QPushButton("Load Preset Profile")
        load_preset_btn.setStyleSheet("QPushButton { padding: 10px 20px; font-size: 14px; background-color: #9C27B0; color: white; border: none; border-radius: 4px; } QPushButton:hover { background-color: #7B1FA2; }")
        load_preset_btn.clicked.connect(self.load_preset_profile)
        header_layout.addWidget(load_preset_btn)

        # Save Profile button (for remapped controls)
        self.save_profile_btn = QPushButton("💾 Save Profile")
        self.save_profile_btn.setStyleSheet("QPushButton { padding: 10px 20px; font-size: 14px; background-color: #2196F3; color: white; border: none; border-radius: 4px; } QPushButton:hover { background-color: #0b7dda; }")
        self.save_profile_btn.clicked.connect(self.save_profile)
        self.save_profile_btn.setEnabled(False)  # Disabled until profile is modified
        header_layout.addWidget(self.save_profile_btn)

        # Export buttons
        self.export_csv_btn = QPushButton("Export CSV")
        self.export_csv_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_csv_btn.clicked.connect(self.export_csv)
        self.export_csv_btn.setEnabled(False)
        header_layout.addWidget(self.export_csv_btn)

        self.export_pdf_btn = QPushButton("Export PDF")
        self.export_pdf_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_pdf_btn.clicked.connect(self.export_pdf)
        self.export_pdf_btn.setEnabled(False)
        header_layout.addWidget(self.export_pdf_btn)

        self.export_word_btn = QPushButton("Export Word")
        self.export_word_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_word_btn.clicked.connect(self.export_word)
        self.export_word_btn.setEnabled(False)
        header_layout.addWidget(self.export_word_btn)

        self.export_graphic_btn = QPushButton("Export Graphic")
        self.export_graphic_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_graphic_btn.clicked.connect(self.export_graphic)
        self.export_graphic_btn.setEnabled(False)
        header_layout.addWidget(self.export_graphic_btn)

        # Help button (rightmost)
        help_btn = QPushButton("Help")
        help_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        help_btn.clicked.connect(self.show_help)
        header_layout.addWidget(help_btn)

        main_layout.addLayout(header_layout)

        # Profile name section with editable field
        profile_name_layout = QHBoxLayout()
        profile_name_layout.setContentsMargins(5, 5, 5, 5)

        profile_name_layout.addWidget(QLabel("Profile Name:"))
        self.profile_name_input = QLineEdit()
        self.profile_name_input.setPlaceholderText("Enter profile name...")
        self.profile_name_input.setEnabled(False)
        profile_name_layout.addWidget(self.profile_name_input, 1)

        self.update_profile_name_btn = QPushButton("Update Name")
        self.update_profile_name_btn.setEnabled(False)
        self.update_profile_name_btn.clicked.connect(self.on_update_profile_name)
        profile_name_layout.addWidget(self.update_profile_name_btn)

        main_layout.addLayout(profile_name_layout)

        # Profile info label (for other info)
        self.profile_info_label = QLabel("No profile loaded")
        self.profile_info_label.setStyleSheet("font-size: 12px; margin: 5px; color: #666;")
        main_layout.addWidget(self.profile_info_label)

        # Filter toolbar
        filter_group = QGroupBox("Filters")
        filter_layout = QHBoxLayout()
        filter_group.setLayout(filter_layout)

        # Search box
        filter_layout.addWidget(QLabel("Search:"))
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Type to filter actions, inputs, or devices...")
        self.search_box.textChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.search_box, 2)

        # Device filter
        filter_layout.addWidget(QLabel("Device:"))
        self.device_filter = QComboBox()
        self.device_filter.addItem("All Devices")
        self.device_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.device_filter, 1)

        # Action Category filter
        filter_layout.addWidget(QLabel("Action Category:"))
        self.actionmap_filter = QComboBox()
        self.actionmap_filter.addItem("All Action Categories")
        self.actionmap_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.actionmap_filter, 1)

        # Key filter
        filter_layout.addWidget(QLabel("Key:"))
        self.key_filter = QComboBox()
        self.key_filter.addItem("All Keys")
        self.key_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.key_filter, 1)

        # Detect Key button
        self.detect_key_btn = QPushButton("Detect Key")
        self.detect_key_btn.setMaximumWidth(100)
        self.detect_key_btn.clicked.connect(self.on_detect_key_clicked)
        filter_layout.addWidget(self.detect_key_btn)

        # Hide unmapped keys checkbox
        self.hide_unmapped_checkbox = QCheckBox("Hide Unmapped Keys")
        self.hide_unmapped_checkbox.stateChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.hide_unmapped_checkbox)

        # Hide mapped keys checkbox
        self.hide_mapped_checkbox = QCheckBox("Hide Mapped Keys")
        self.hide_mapped_checkbox.stateChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.hide_mapped_checkbox)

        # Show detailed checkbox
        self.show_detailed_checkbox = QCheckBox("Show Detailed")
        self.show_detailed_checkbox.stateChanged.connect(self.toggle_detailed_view)
        filter_layout.addWidget(self.show_detailed_checkbox)

        # Clear filters button
        clear_btn = QPushButton("Clear Filters")
        clear_btn.clicked.connect(self.clear_filters)
        filter_layout.addWidget(clear_btn)

        main_layout.addWidget(filter_group)

        # Create tab widget for different views
        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)

        # Tab 1: Controls Table View
        table_tab = QWidget()
        table_layout = QVBoxLayout()
        table_tab.setLayout(table_layout)

        # Splitter for info and table
        splitter = QSplitter(Qt.Orientation.Vertical)

        # Profile summary text
        self.profile_summary = QTextEdit()
        self.profile_summary.setReadOnly(True)
        self.profile_summary.setMaximumHeight(150)
        self.profile_summary.setPlaceholderText("Profile summary will appear here...")
        splitter.addWidget(self.profile_summary)

        # Controls table
        self.controls_table = QTableWidget()
        self.controls_table.setColumnCount(8)  # Max columns for detailed view
        self.controls_table.setHorizontalHeaderLabels([
            "Action Category", "Action", "Label", "Input", "Input Code", "Activation", "Device", "Edit"
        ])
        self.controls_table.horizontalHeader().setStretchLastSection(False)
        self.controls_table.setAlternatingRowColors(True)
        self.controls_table.setSortingEnabled(True)
        self.controls_table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        self.controls_table.itemChanged.connect(self.on_cell_edited)
        self.controls_table.itemDoubleClicked.connect(self.on_item_double_clicked)

        # Enable context menu for remapping
        self.controls_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.controls_table.customContextMenuRequested.connect(self.show_table_context_menu)

        # Set custom delegate for column 2 (Label) to handle text selection
        delegate = SelectAllDelegate(self.controls_table, self)
        self.controls_table.setItemDelegateForColumn(2, delegate)

        splitter.addWidget(self.controls_table)

        # Track editing state
        self._editing_item = None
        self._editing_default_text = None  # Stores the default label (without custom override)

        # Track editable columns (column 2 is Label, editable in both views)
        self.editable_columns = {2}  # Column 2: Label

        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        table_layout.addWidget(splitter)
        self.tab_widget.addTab(table_tab, "Controls Table")

        # Tab 2: Device View (PDF-based device visualization) - Optional
        templates_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "visual-templates")
        if PDF_WIDGET_AVAILABLE:
            self.pdf_device_widget = PDF_WIDGET_CLASS(templates_dir)
            self.pdf_device_widget.export_available_changed.connect(self.export_graphic_btn.setEnabled)
            self.tab_widget.addTab(self.pdf_device_widget, "Device View")
        else:
            self.pdf_device_widget = None
            logger.warning("Device View tab disabled - no PDF widget available")

        # Tab 3: Config (Device Configuration)
        from gui.config_tab import ConfigTab
        self.config_tab = ConfigTab()
        self.tab_widget.addTab(self.config_tab, "Config")

        # Connect config tab device updates to Device View
        if self.pdf_device_widget:
            self.config_tab.devices_changed.connect(self.pdf_device_widget.set_connected_devices)
            # Pass initial devices to Device View
            self.pdf_device_widget.set_connected_devices(self.config_tab.current_devices)

        # Tab 4: About
        about_tab = QWidget()
        about_layout = QVBoxLayout()
        about_tab.setLayout(about_layout)

        # About content in a scrollable text browser
        about_browser = QTextBrowser()
        about_browser.setOpenExternalLinks(True)

        try:
            # Load ABOUT.md file
            about_path = get_resource_path("ABOUT.md")
            with open(about_path, 'r', encoding='utf-8') as f:
                about_content = f.read()

            # Add version to the first heading
            about_content = about_content.replace(
                "# Star Citizen Profile Editor",
                f"# Star Citizen Profile Editor v{self.version}"
            )

            # Convert markdown to HTML
            about_html = self.markdown_to_html(about_content)
            about_browser.setHtml(about_html)
        except Exception as e:
            logger.error(f"Error loading ABOUT.md: {e}", exc_info=True)
            about_browser.setHtml("<h1>About</h1><p>Unable to load about information.</p>")
        about_layout.addWidget(about_browser)

        self.tab_widget.addTab(about_tab, "About")

        # Connect tab change signal to sync device selection
        self.tab_widget.currentChanged.connect(self.on_tab_changed)

        # Footer with Osiris DevWorks (left) and PayPal (right)
        footer_layout = QHBoxLayout()

        # Osiris DevWorks button (left side)
        self.osiris_button = QLabel()
        osiris_image_path = get_resource_path(os.path.join("assets", "osiris-devworks.png"))

        # Try to load Osiris image, fall back to text if not found
        if os.path.exists(osiris_image_path):
            pixmap = QPixmap(osiris_image_path)
            # Scale to reasonable size (max height 40px to accommodate the logo design)
            if pixmap.height() > 40:
                pixmap = pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.osiris_button.setPixmap(pixmap)
        else:
            # Fallback to styled text button
            self.osiris_button.setText("Osiris DevWorks")
            self.osiris_button.setStyleSheet("""
                QLabel {
                    background-color: #1a1f2e;
                    color: #c9a961;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #242938;
                }
            """)

        self.osiris_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.osiris_button.mousePressEvent = self.open_discord_link
        footer_layout.addWidget(self.osiris_button)

        # Stretch to push donation buttons to the right
        footer_layout.addStretch()

        # Donation label
        donation_label = QLabel("Support this project:")
        donation_label.setStyleSheet("font-size: 12px; color: #666; margin-right: 5px;")
        footer_layout.addWidget(donation_label)

        # PayPal button (right side)
        self.paypal_button = QLabel()
        paypal_image_path = get_resource_path(os.path.join("assets", "paypal.png"))

        # Try to load PayPal image, fall back to text if not found
        if os.path.exists(paypal_image_path):
            paypal_pixmap = QPixmap(paypal_image_path)
            # Scale to match Osiris button (max height 40px)
            if paypal_pixmap.height() > 40:
                paypal_pixmap = paypal_pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.paypal_button.setPixmap(paypal_pixmap)
        else:
            # Fallback to styled text button
            self.paypal_button.setText("Donate via PayPal")
            self.paypal_button.setStyleSheet("""
                QLabel {
                    background-color: #0070ba;
                    color: white;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #005ea6;
                }
            """)

        self.paypal_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.paypal_button.mousePressEvent = self.open_paypal_donation
        footer_layout.addWidget(self.paypal_button)

        # Spacer between PayPal and Venmo
        footer_layout.addSpacing(10)

        # Venmo button (right side)
        self.venmo_button = QLabel()
        venmo_image_path = get_resource_path(os.path.join("assets", "venmo.png"))

        # Try to load Venmo image, fall back to text button
        if os.path.exists(venmo_image_path):
            venmo_pixmap = QPixmap(venmo_image_path)
            # Scale to match Osiris button (max height 40px)
            if venmo_pixmap.height() > 40:
                venmo_pixmap = venmo_pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.venmo_button.setPixmap(venmo_pixmap)
        else:
            # Fallback to styled text button
            self.venmo_button.setText("Venmo")
            self.venmo_button.setStyleSheet("""
                QLabel {
                    background-color: #008CFF;
                    color: white;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #0074D9;
                }
            """)

        self.venmo_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.venmo_button.mousePressEvent = self.open_venmo_donation
        footer_layout.addWidget(self.venmo_button)

        main_layout.addLayout(footer_layout)

        # Status bar
        self.statusBar().showMessage("Ready")

    def restore_window_state(self):
        """Restore saved window geometry and state"""
        geometry = self.settings.get_window_geometry()
        if geometry:
            self.restoreGeometry(geometry)
            logger.debug("Restored window geometry")

        state = self.settings.get_window_state()
        if state:
            self.restoreState(state)
            logger.debug("Restored window state")

    def setup_tray_icon(self):
        """Set up the system tray icon and its context menu"""
        # Load application icon
        icon_path = get_resource_path("assets/icon.ico")
        icon = QIcon(icon_path)

        # Create system tray icon
        self.tray_icon = QSystemTrayIcon(icon, self)

        # Create tray icon context menu
        tray_menu = QMenu()

        # Show/Restore action
        show_action = QAction("Show Star Citizen Profile Editor", self)
        show_action.triggered.connect(self.restore_from_tray)
        tray_menu.addAction(show_action)

        tray_menu.addSeparator()

        # Exit action
        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.quit_application)
        tray_menu.addAction(exit_action)

        # Set menu and connect activation signal
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self.on_tray_icon_activated)

        # Set tooltip
        self.tray_icon.setToolTip("Star Citizen Profile Editor")

        # Show tray icon immediately
        self.tray_icon.show()

        logger.info("System tray icon initialized and shown")

    def on_tray_icon_activated(self, reason):
        """
        Handle tray icon activation (click)

        Args:
            reason: QSystemTrayIcon.ActivationReason enum value
        """
        # Restore window on left click or double click
        if reason in (QSystemTrayIcon.ActivationReason.Trigger,
                      QSystemTrayIcon.ActivationReason.DoubleClick):
            self.restore_from_tray()

    def restore_from_tray(self):
        """Restore the window from system tray"""
        self.show()
        self.setWindowState(self.windowState() & ~Qt.WindowState.WindowMinimized | Qt.WindowState.WindowActive)
        self.activateWindow()
        self.raise_()
        logger.debug("Window restored from system tray")

    def quit_application(self):
        """Quit the application cleanly"""
        logger.info("Application quit requested from tray icon")
        self.is_closing = True
        if self.tray_icon:
            self.tray_icon.hide()
        self.close()

    def changeEvent(self, event):
        """Handle window state changes (minimize, maximize, etc.)"""
        if event.type() == QEvent.Type.WindowStateChange:
            # Check if window is being minimized
            if self.windowState() & Qt.WindowState.WindowMinimized:
                # Check if minimize-to-tray is enabled
                if self.settings.get_minimize_to_tray_enabled():
                    # Hide to tray instead of showing in taskbar
                    QTimer.singleShot(0, self.hide)
                    event.ignore()
                    logger.debug("Window minimized to system tray")

                    # Show notification on first minimize
                    if self.tray_icon and not hasattr(self, '_tray_notification_shown'):
                        self.tray_icon.showMessage(
                            "Star Citizen Profile Editor",
                            "Application minimized to system tray",
                            QSystemTrayIcon.MessageIcon.Information,
                            2000  # 2 seconds
                        )
                        self._tray_notification_shown = True
                    return

        # Call parent implementation for other events
        super().changeEvent(event)

    def closeEvent(self, event):
        """Save window state before closing"""
        self.settings.set_window_geometry(self.saveGeometry())
        self.settings.set_window_state(self.saveState())
        logger.debug("Saved window state")

        # Hide and clean up tray icon when closing
        if self.tray_icon:
            self.tray_icon.hide()

        event.accept()

    def auto_load_last_profile(self):
        """Automatically load the last opened profile if it exists"""
        last_profile_path = self.settings.get_last_profile_path()

        if not last_profile_path:
            logger.info("No last profile found")
            return

        if not os.path.exists(last_profile_path):
            logger.warning(f"Last profile no longer exists: {last_profile_path}")
            self.settings.clear_last_profile_path()
            return

        logger.info(f"Auto-loading last profile: {last_profile_path}")
        self.load_profile_from_path(last_profile_path)

    def import_profile(self):
        """Handle profile import button click"""
        # Determine starting directory with preference order:
        # 1. Star Citizen profiles directory (if configured and exists)
        # 2. Last profile directory (if available)
        # 3. Empty (let dialog use default)
        start_dir = ""

        sc_profiles_dir = self.settings.get_sc_profiles_directory()
        if sc_profiles_dir and os.path.exists(sc_profiles_dir):
            start_dir = sc_profiles_dir
        else:
            last_path = self.settings.get_last_profile_path()
            if last_path and os.path.exists(last_path):
                start_dir = os.path.dirname(last_path)

        # Open file dialog
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Star Citizen Profile XML",
            start_dir,
            "XML Files (*.xml);;All Files (*)"
        )

        if not file_path:
            return  # User cancelled

        self.load_profile_from_path(file_path)

    def load_profile_from_path(self, file_path: str):
        """
        Load a profile from the given file path

        Args:
            file_path: Path to the profile XML file
        """
        try:
            # Parse the profile with overlay system (always enabled)
            self.statusBar().showMessage(f"Loading profile: {file_path}")
            logger.info(f"Loading profile: {file_path}")

            # Always use overlay system with blank.xml as master template
            parser = ProfileParser(file_path, use_bundled_defaults=True)
            self.current_profile = parser.parse()
            self.current_profile_path = file_path

            # Save as last opened profile
            self.settings.set_last_profile_path(file_path)

            # Update UI with profile data
            self.display_profile()

            # Enable export buttons
            self.export_csv_btn.setEnabled(True)
            self.export_pdf_btn.setEnabled(True)
            self.export_word_btn.setEnabled(True)

            # Show status
            status_msg = f"Successfully loaded: {self.current_profile.profile_name}"
            self.statusBar().showMessage(status_msg)
            logger.info(f"Successfully loaded profile: {self.current_profile.profile_name}")

        except FileNotFoundError:
            QMessageBox.critical(self, "Error", f"File not found: {file_path}")
            self.statusBar().showMessage("Error: File not found")
            logger.error(f"File not found: {file_path}")
        except ValueError as e:
            QMessageBox.critical(self, "Parse Error", f"Failed to parse XML file:\n{str(e)}")
            self.statusBar().showMessage("Error: Failed to parse XML")
            logger.error(f"Parse error: {e}", exc_info=True)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An unexpected error occurred:\n{str(e)}")
            self.statusBar().showMessage("Error loading profile")
            logger.error(f"Error loading profile: {e}", exc_info=True)

    def create_new_profile(self):
        """Create a new profile from blank.xml without file dialog"""
        try:
            # Find blank.xml in example-profiles directory
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            blank_profile_path = os.path.join(project_root, "example-profiles", "blank.xml")

            if not os.path.exists(blank_profile_path):
                QMessageBox.critical(
                    self,
                    "Error",
                    f"Blank profile not found at {blank_profile_path}\n\n"
                    "Please ensure blank.xml exists in the example-profiles directory."
                )
                logger.error(f"Blank profile not found: {blank_profile_path}")
                return

            # Load the blank profile
            self.load_profile_from_path(blank_profile_path)

            # Add currently connected joystick devices to the new profile
            self._add_missing_joystick_devices()

            # Refresh UI to show the newly added devices
            self.display_profile()

            # Refresh Device View widget with updated devices
            if self.pdf_device_widget:
                self.pdf_device_widget.load_profile(self.current_profile)
                logger.debug("Refreshed Device View with newly added joystick devices")

            # Set a default filename for save
            from datetime import datetime
            default_name = f"new_profile_{datetime.now().strftime('%Y%m%d')}.xml"
            self.current_profile_path = None  # Force "Save As" dialog on first save
            self.statusBar().showMessage(f"New profile created. Use 'Save Profile' to save with filename like '{default_name}'")

            logger.info("New profile created from blank.xml")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create new profile:\n{str(e)}")
            logger.error(f"Error creating new profile: {e}", exc_info=True)

    def load_preset_profile(self):
        """Show dialog to select and load a preset profile"""
        try:
            # Get bundled presets directory
            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

            presets_dir = os.path.join(base_path, 'default-bindings', 'presets')

            if not os.path.exists(presets_dir):
                QMessageBox.warning(
                    self,
                    "Presets Not Found",
                    "Preset profiles directory not found. Please reinstall the application."
                )
                logger.warning(f"Presets directory not found: {presets_dir}")
                return

            # Get list of preset files
            try:
                preset_files = sorted([f for f in os.listdir(presets_dir) if f.endswith('.xml')])
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "Error Reading Presets",
                    f"Failed to read preset profiles directory:\n{str(e)}"
                )
                logger.error(f"Error reading presets directory: {e}", exc_info=True)
                return

            if not preset_files:
                QMessageBox.information(
                    self,
                    "No Presets",
                    "No preset profiles found."
                )
                return

            # Create selection dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Load Preset Profile")
            dialog.setMinimumWidth(500)

            layout = QVBoxLayout()

            # Instructions
            label = QLabel("Select a preset profile to load as your starting point:")
            layout.addWidget(label)

            # List of presets with friendly names
            preset_list = QListWidget()
            preset_mapping = {}

            for filename in preset_files:
                # Convert filename to friendly name
                friendly_name = filename.replace('layout_', '').replace('.xml', '').replace('_', ' ').title()
                preset_list.addItem(friendly_name)
                preset_mapping[friendly_name] = filename

            preset_list.itemDoubleClicked.connect(lambda: dialog.accept())
            layout.addWidget(preset_list)

            # Dialog buttons
            buttons = QHBoxLayout()
            ok_btn = QPushButton("Load")
            cancel_btn = QPushButton("Cancel")
            ok_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            buttons.addStretch()
            buttons.addWidget(ok_btn)
            buttons.addWidget(cancel_btn)
            layout.addLayout(buttons)

            dialog.setLayout(layout)

            # Show dialog
            if dialog.exec() == QDialog.DialogCode.Accepted:
                selected_items = preset_list.selectedItems()
                if selected_items:
                    friendly_name = selected_items[0].text()
                    filename = preset_mapping[friendly_name]
                    preset_path = os.path.join(presets_dir, filename)

                    # Load the preset profile
                    self.load_profile_from_path(preset_path)

                    # Ensure currently connected joystick devices are in the profile
                    self._add_missing_joystick_devices()

                    # Refresh UI to show the newly added devices
                    self.display_profile()

                    # Refresh Device View widget with updated devices
                    if self.pdf_device_widget:
                        self.pdf_device_widget.load_profile(self.current_profile)
                        logger.debug("Refreshed Device View with connected joystick devices")

                    # Mark as modified since it's a starting point for customization
                    self.current_profile.mark_modified()
                    self.save_profile_btn.setEnabled(True)

                    # Update window title to indicate it's a preset
                    self.setWindowTitle(f"Star Citizen Profile Editor v{self.version} - {friendly_name} (Preset - Modified)")

                    self.statusBar().showMessage(f"Loaded preset: {friendly_name}", 3000)
                    logger.info(f"Loaded preset profile: {friendly_name}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load preset profile:\n{str(e)}")
            logger.error(f"Error loading preset profile: {e}", exc_info=True)

    def display_profile(self):
        """Display the loaded profile data"""
        if not self.current_profile:
            return

        profile = self.current_profile

        # Update profile name input field and enable editing
        self.profile_name_input.setText(profile.profile_name)
        self.profile_name_input.setEnabled(True)
        self.update_profile_name_btn.setEnabled(True)

        # Update profile info label
        self.profile_info_label.setText(f"Profile: {profile.profile_name}")

        # Check if any joystick devices in profile are disconnected
        self._check_and_warn_disconnected_devices(profile)

        # Build summary text
        summary_parts = []
        summary_parts.append(f"Profile Name: {profile.profile_name}")
        summary_parts.append(f"Devices: {len(profile.devices)}")

        # List devices
        for device in profile.devices:
            device_name = device.product_name if device.product_name else device.device_type.capitalize()
            summary_parts.append(f"  - {device.device_type.capitalize()} {device.instance}: {device_name}")

        summary_parts.append(f"\nAction Categories: {len(profile.action_maps)}")
        summary_parts.append(f"Total Bindings: {len(profile.get_all_bindings())}")

        if profile.categories:
            summary_parts.append(f"\nCategories: {', '.join(profile.categories[:5])}")
            if len(profile.categories) > 5:
                summary_parts.append(f"  ... and {len(profile.categories) - 5} more")

        self.profile_summary.setText('\n'.join(summary_parts))

        # Populate controls table
        self.populate_controls_table()

        # Apply filters to show the rows (important when loading a new profile)
        self.apply_filters()

        # Load profile into Device View widget
        if self.pdf_device_widget:
            self.pdf_device_widget.load_profile(profile)

    def populate_controls_table(self):
        """Populate the controls table with bindings"""
        if not self.current_profile:
            return

        # Disable sorting while populating
        self.controls_table.setSortingEnabled(False)

        # Block signals to prevent itemChanged from triggering during population
        self.controls_table.blockSignals(True)

        # Get all bindings and store for filtering
        self.all_bindings = []
        for action_map in self.current_profile.action_maps:
            for binding in action_map.actions:
                self.all_bindings.append((action_map.name, binding))

        # Populate filter dropdowns
        self.populate_filter_dropdowns()

        # Set row count
        self.controls_table.setRowCount(len(self.all_bindings))

        # Determine if we're in detailed view
        is_detailed = self.show_detailed_checkbox.isChecked()

        # Configure column visibility based on view mode
        if is_detailed:
            # Detailed view: Show columns 0, 1, 2, 3, 4, 5, 6, 7
            # 0: Action Category, 1: Action (original), 2: Label, 3: Input, 4: Input Code, 5: Activation, 6: Device, 7: Edit
            self.controls_table.setColumnHidden(1, False)
            self.controls_table.setColumnHidden(4, False)  # Show raw input code column
            self.controls_table.setColumnHidden(5, False)  # Show activation mode column
            self.controls_table.setColumnHidden(7, False)  # Show edit button
        else:
            # Default view: Show columns 0, 2, 3, 6, 7
            # 0: Action Category, 2: Label, 3: Input, 6: Device, 7: Edit
            self.controls_table.setColumnHidden(1, True)   # Hide original Action
            self.controls_table.setColumnHidden(4, True)   # Hide raw input code column
            self.controls_table.setColumnHidden(5, True)   # Hide activation mode column
            self.controls_table.setColumnHidden(6, False)  # Show device column

        # Populate table
        for row, (action_map_name, binding) in enumerate(self.all_bindings):
            # Column 0: Action Category
            action_map_label = LabelGenerator.generate_actionmap_label(action_map_name)
            action_map_item = QTableWidgetItem(action_map_label)
            action_map_item.setFlags(action_map_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 0, action_map_item)

            # Column 1: Action (original auto-generated) - only visible in detailed view
            action_label_original = LabelGenerator.generate_action_label(binding.action_name)
            action_original_item = QTableWidgetItem(action_label_original)
            action_original_item.setFlags(action_original_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 1, action_original_item)

            # Column 2: Label (editable custom label) - visible in both views
            action_label_override = LabelGenerator.get_action_label(binding.action_name, binding)
            label_item = QTableWidgetItem(action_label_override)
            # Store binding reference in the item for later retrieval during editing
            label_item.setData(Qt.ItemDataRole.UserRole, (action_map_name, binding))
            self.controls_table.setItem(row, 2, label_item)

            # Check if this is an unmapped action (input code ends with underscore)
            is_unmapped = binding.input_code.rstrip().endswith('_')

            # Column 3: Input (human-readable with raw code in tooltip)
            if is_unmapped:
                # For unmapped actions, show blank
                input_item = QTableWidgetItem("")
            else:
                input_label = LabelGenerator.generate_input_label(binding.input_code)
                input_item = QTableWidgetItem(input_label)
                input_item.setToolTip(binding.input_code)  # Show raw code in tooltip
            input_item.setFlags(input_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 3, input_item)

            # Column 4: Raw input code (shown in detailed view, right after Input)
            input_code_item = QTableWidgetItem(binding.input_code)
            input_code_item.setFlags(input_code_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 4, input_code_item)

            # Column 5: Activation Mode (only visible in detailed view)
            # Default to "press" if no activation mode is specified
            activation_text = binding.activation_mode if binding.activation_mode else "press"
            activation_item = QTableWidgetItem(activation_text)
            activation_item.setFlags(activation_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 5, activation_item)

            # Column 6: Device (parsed from input code)
            # Only show device for mapped actions; unmapped actions should be blank
            if is_unmapped:
                device_item = QTableWidgetItem("")
            else:
                device = self.parse_device_from_input(binding.input_code)
                device_item = QTableWidgetItem(device)
            device_item.setFlags(device_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 6, device_item)

            # Column 7: Edit button
            edit_btn = QPushButton("✏️")
            edit_btn.setMaximumWidth(40)
            edit_btn.clicked.connect(lambda checked, r=row: self.on_edit_input_clicked(r))
            self.controls_table.setCellWidget(row, 7, edit_btn)

        # Re-enable signals
        self.controls_table.blockSignals(False)

        # Re-enable sorting
        self.controls_table.setSortingEnabled(True)

        # Resize columns to content
        self.controls_table.resizeColumnsToContents()

        # Set column widths based on view mode
        if is_detailed:
            self.controls_table.setColumnWidth(0, 150)  # Action Category
            self.controls_table.setColumnWidth(1, 200)  # Action (Original)
            self.controls_table.setColumnWidth(2, 200)  # Label
            self.controls_table.setColumnWidth(3, 200)  # Input
            self.controls_table.setColumnWidth(4, 150)  # Input Code (raw)
            self.controls_table.setColumnWidth(5, 100)  # Activation Mode
            self.controls_table.setColumnWidth(6, 120)  # Device
            self.controls_table.setColumnWidth(7, 50)   # Edit button
        else:
            self.controls_table.setColumnWidth(0, 150)  # Action Category
            self.controls_table.setColumnWidth(2, 200)  # Label
            self.controls_table.setColumnWidth(3, 200)  # Input
            self.controls_table.setColumnWidth(6, 120)  # Device
            self.controls_table.setColumnWidth(7, 50)   # Edit button

    def parse_device_from_input(self, input_code: str) -> str:
        """Parse device type from input code using actual connected device info"""
        if input_code.startswith('kb'):
            return "Keyboard"
        elif input_code.startswith('js'):
            # Extract joystick instance
            import re
            match = re.match(r'js(\d+)_', input_code)
            if match:
                instance = int(match.group(1))

                # First, try to find the actual connected device with this instance number
                if hasattr(self, 'config_tab') and self.config_tab and hasattr(self.config_tab, 'current_devices'):
                    for connected_device in self.config_tab.current_devices:
                        if (connected_device.get('type') == 'joystick' and
                            connected_device.get('instance') == instance):
                            device_name = connected_device.get('name', f"Joystick {instance}")
                            # Split composite devices (e.g., VKB Gladiator + SEM)
                            return get_device_for_input(device_name, input_code)

                # Fallback: try to find device in profile (for backwards compatibility)
                for device in self.current_profile.devices:
                    if device.device_type == 'joystick' and device.instance == instance:
                        device_name = device.product_name if device.product_name else f"Joystick {instance}"
                        # Split composite devices (e.g., VKB Gladiator + SEM)
                        return get_device_for_input(device_name, input_code)

                return f"Joystick {instance}"
        elif 'mouse' in input_code.lower():
            return "Mouse"
        return "Unknown"

    def populate_filter_dropdowns(self):
        """Populate the device and action map filter dropdowns"""
        # Get unique devices
        devices = set()
        action_maps = set()
        input_codes = set()  # Collect unique input codes for key filter

        for action_map_name, binding in self.all_bindings:
            device = self.parse_device_from_input(binding.input_code)
            devices.add(device)
            action_map_label = LabelGenerator.generate_actionmap_label(action_map_name)
            action_maps.add(action_map_label)

            # Collect mapped input codes (exclude unmapped ones ending with '_')
            input_code = binding.input_code.strip()
            if input_code and not input_code.endswith('_'):
                input_codes.add(input_code)

        # Always include keyboard, mouse, and joystick options
        devices.add("Keyboard")
        devices.add("Mouse")
        # Joysticks will be added dynamically if present in bindings

        # Save current selections before clearing
        current_device = self.device_filter.currentText()
        current_actionmap = self.actionmap_filter.currentText()
        current_key = self.key_filter.currentText()

        # Update device filter
        self.device_filter.blockSignals(True)
        self.device_filter.clear()
        self.device_filter.addItem("All Devices")
        for device in sorted(devices):
            self.device_filter.addItem(device)

        # Restore previous selection if it still exists
        if current_device:
            index = self.device_filter.findText(current_device)
            if index >= 0:
                self.device_filter.setCurrentIndex(index)

        self.device_filter.blockSignals(False)

        # Update action category filter
        self.actionmap_filter.blockSignals(True)
        self.actionmap_filter.clear()
        self.actionmap_filter.addItem("All Action Categories")
        for action_map in sorted(action_maps):
            self.actionmap_filter.addItem(action_map)

        # Restore previous selection if it still exists
        if current_actionmap:
            index = self.actionmap_filter.findText(current_actionmap)
            if index >= 0:
                self.actionmap_filter.setCurrentIndex(index)

        self.actionmap_filter.blockSignals(False)

        # Update key filter
        self.key_filter.blockSignals(True)
        self.key_filter.clear()
        self.key_filter.addItem("All Keys")

        # Group input codes by device type: keyboard, mouse, joystick
        keyboard_codes = sorted([ic for ic in input_codes if ic.startswith('kb')])
        mouse_codes = sorted([ic for ic in input_codes if ic.startswith('mouse')])
        joystick_codes = sorted([ic for ic in input_codes if ic.startswith('js')])

        # Add in order: keyboard, mouse, joystick
        for input_code in keyboard_codes + mouse_codes + joystick_codes:
            # Generate human-readable label
            label = LabelGenerator.generate_input_label(input_code)
            # Format: "Label (raw_code)" e.g., "Keyboard O (kb1_o)"
            display_text = f"{label} ({input_code})"
            # Store raw input code as user data for filtering
            self.key_filter.addItem(display_text, input_code)

        # Restore previous selection if it still exists
        if current_key:
            index = self.key_filter.findText(current_key)
            if index >= 0:
                self.key_filter.setCurrentIndex(index)

        self.key_filter.blockSignals(False)

    def apply_filters(self):
        """Apply current filter settings to the table"""
        if not self.all_bindings:
            return

        search_text = self.search_box.text().lower()
        device_filter = self.device_filter.currentText()
        actionmap_filter = self.actionmap_filter.currentText()
        key_filter = self.key_filter.currentText()
        # Extract raw input code from dropdown (stored as user data)
        key_filter_code = None
        if key_filter != "All Keys":
            current_index = self.key_filter.currentIndex()
            key_filter_code = self.key_filter.itemData(current_index)
        hide_unmapped = self.hide_unmapped_checkbox.isChecked()
        hide_mapped = self.hide_mapped_checkbox.isChecked()

        # Show/hide rows based on filters
        for row in range(self.controls_table.rowCount()):
            show_row = True

            # Search filter - check all columns except Edit button (5) and hidden raw input code (6)
            if search_text:
                row_text = ""
                for col in [0, 1, 2, 3, 4]:  # Skip columns 5 (Edit button) and 6 (hidden input code)
                    item = self.controls_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "

                if search_text not in row_text:
                    show_row = False

            # Device filter
            if show_row and device_filter != "All Devices":
                device_item = self.controls_table.item(row, 6)  # Device is column 6
                if device_item:
                    device_text = device_item.text().strip()
                    # Only hide row if it has a device and the device doesn't match the filter
                    # (Don't hide rows with empty device - those are unmapped actions)
                    if device_text and device_text != device_filter:
                        show_row = False
                else:
                    # No device item means unmapped action - don't hide it for device filter
                    pass

            # Action category filter
            if show_row and actionmap_filter != "All Action Categories":
                actionmap_item = self.controls_table.item(row, 0)
                if actionmap_item and actionmap_item.text() != actionmap_filter:
                    show_row = False

            # Key filter - check if row's input code matches selected key
            if show_row and key_filter != "All Keys" and key_filter_code:
                input_code_item = self.controls_table.item(row, 4)  # Raw input code is column 4
                if input_code_item:
                    row_input_code = input_code_item.text().strip()
                    # Only match exact input codes
                    if row_input_code != key_filter_code:
                        show_row = False

            # Unmapped keys filter - check if input code ends with underscore
            # Unmapped keys are stored as "js1_ ", "kb1_", "mouse_", etc. (ending with underscore+space or just underscore)
            # Column 4 contains the raw input code
            if show_row and hide_unmapped:
                input_code_item = self.controls_table.item(row, 4)  # Raw input code is in column 4
                if input_code_item:
                    input_code = input_code_item.text().strip()
                    # Hide rows where input code is empty or ends with underscore (unmapped)
                    if not input_code or input_code.endswith('_'):
                        show_row = False

            # Mapped keys filter - opposite of unmapped: hide rows with valid input codes
            # Hides rows where input code is NOT empty and does NOT end with underscore (mapped)
            if show_row and hide_mapped:
                input_code_item = self.controls_table.item(row, 4)  # Raw input code is in column 4
                if input_code_item:
                    input_code = input_code_item.text().strip()
                    # Hide rows where input code is mapped (not empty and doesn't end with underscore)
                    if input_code and not input_code.endswith('_'):
                        show_row = False

            self.controls_table.setRowHidden(row, not show_row)

        # Update status bar
        visible_rows = sum(1 for row in range(self.controls_table.rowCount())
                          if not self.controls_table.isRowHidden(row))
        total_rows = self.controls_table.rowCount()
        self.statusBar().showMessage(f"Showing {visible_rows} of {total_rows} bindings")

    def toggle_detailed_view(self):
        """Toggle between default and detailed view"""
        if not self.current_profile:
            return

        # Check if user is currently editing a cell
        current_item = self.controls_table.currentItem()
        if current_item and current_item.column() == 2:  # Label column
            # Check if there's an active editor widget
            editor = self.controls_table.cellWidget(current_item.row(), current_item.column())
            if not editor:
                # Check if the item is in edit mode via QTableWidget's state
                state = self.controls_table.state()
                from PyQt6.QtWidgets import QAbstractItemView
                if state == QAbstractItemView.State.EditingState:
                    # There's an active edit - commit it by setting focus elsewhere
                    # This will trigger the itemChanged signal with the current editor value
                    self.controls_table.setCurrentItem(None)
                    # Small delay to let the edit complete
                    from PyQt6.QtCore import QTimer
                    QTimer.singleShot(0, self._complete_toggle_view)
                    return

        # No active edit, proceed immediately
        self._complete_toggle_view()

    def _complete_toggle_view(self):
        """Complete the view toggle after any pending edits are committed"""
        # Repopulate the table with the new view mode
        self.populate_controls_table()

        # Reapply filters
        self.apply_filters()

        # Update status message
        view_mode = "detailed" if self.show_detailed_checkbox.isChecked() else "default"
        self.statusBar().showMessage(f"Switched to {view_mode} view")

    def on_detect_key_clicked(self):
        """Handle Detect Key button click for key filter"""
        if self.detecting_input_for_filter:
            # Cancel ongoing detection
            if self.input_detector:
                self.input_detector.stop_detection()
            self.detecting_input_for_filter = False
            self.detect_key_btn.setText("Detect Key")
            self.detect_key_btn.setEnabled(True)
            logger.info("Key filter input detection cancelled")
            return

        # Start detection
        try:
            # Lazy initialization of InputDetector
            if not self.input_detector:
                from utils.input_detector import InputDetector
                self.input_detector = InputDetector()

            self.detecting_input_for_filter = True
            self.detect_key_btn.setText("Listening... (Cancel)")
            logger.info("Starting input detection for key filter")

            # Start input detection with 10 second timeout
            detector_thread = self.input_detector.start_detection(timeout_ms=10000)
            detector_thread.input_detected.connect(self.on_key_filter_input_detected)
            detector_thread.detection_cancelled.connect(self.on_key_filter_detection_cancelled)

        except Exception as e:
            logger.error(f"Error starting input detection for key filter: {e}", exc_info=True)
            QMessageBox.critical(
                self,
                "Detection Error",
                f"Failed to start input detection:\n{str(e)}"
            )
            self.detecting_input_for_filter = False
            self.detect_key_btn.setText("Detect Key")

    def on_key_filter_input_detected(self, input_code: str, description: str):
        """Handle detected input for key filter"""
        logger.info(f"Input detected for key filter: {input_code} - {description}")

        # Reset button state
        self.detecting_input_for_filter = False
        self.detect_key_btn.setText("Detect Key")

        # Find the detected input code in the dropdown
        # Search by user data (raw input code)
        for i in range(self.key_filter.count()):
            item_data = self.key_filter.itemData(i)
            if item_data == input_code:
                self.key_filter.setCurrentIndex(i)
                logger.info(f"Set key filter to: {self.key_filter.currentText()}")
                # apply_filters() will be called automatically via currentTextChanged signal
                return

        # Input code not found in filter - might be unmapped or not in current profile
        logger.warning(f"Detected input code '{input_code}' not found in key filter dropdown")
        QMessageBox.information(
            self,
            "Key Not Found",
            f"The detected key ({description}) is not currently mapped in your profile.\n\n"
            f"Only mapped keys appear in the filter dropdown."
        )

    def on_key_filter_detection_cancelled(self):
        """Handle input detection timeout/cancellation for key filter"""
        logger.info("Key filter input detection timed out or cancelled")

        self.detecting_input_for_filter = False
        self.detect_key_btn.setText("Detect Key")

        QMessageBox.information(
            self,
            "No Input Detected",
            "No input was detected within 10 seconds.\n\nPlease try again."
        )

    def clear_filters(self):
        """Clear all filters"""
        self.search_box.clear()
        self.device_filter.setCurrentIndex(0)
        self.actionmap_filter.setCurrentIndex(0)
        self.key_filter.setCurrentIndex(0)  # Reset to "All Keys"

        # Cancel any ongoing input detection for key filter
        if self.detecting_input_for_filter:
            if self.input_detector:
                self.input_detector.stop_detection()
            self.detecting_input_for_filter = False
            self.detect_key_btn.setText("Detect Key")

        self.hide_unmapped_checkbox.setChecked(False)
        self.hide_mapped_checkbox.setChecked(False)
        self.statusBar().showMessage(f"Filters cleared - showing all {self.controls_table.rowCount()} bindings")

    def on_tab_changed(self, index: int):
        """Handle tab change - sync device selection when switching to Device View tab"""
        # Get the currently filtered device from the control table
        filtered_device = self.device_filter.currentText()

        if index == 1 and self.pdf_device_widget:
            # Device View tab
            if filtered_device and filtered_device != "All Devices":
                self.pdf_device_widget.select_device_by_name(filtered_device)

    def show_table_context_menu(self, position):
        """Show context menu for table row"""
        # Get the item at the clicked position
        item = self.controls_table.itemAt(position)
        if not item:
            return

        # Get binding data from the item
        binding_data = item.data(Qt.ItemDataRole.UserRole)
        if not binding_data:
            return

        action_map_name, binding = binding_data

        # Create context menu
        from PyQt6.QtWidgets import QMenu
        menu = QMenu(self)

        # Add "Remap Button..." action
        remap_action = menu.addAction("🎮 Remap Button...")
        remap_action.triggered.connect(lambda: self.show_remap_dialog(binding))

        # Add "Edit Label..." action
        edit_label_action = menu.addAction("✏️ Edit Label...")
        edit_label_action.triggered.connect(lambda: self.edit_label_from_context(binding))

        # Show menu at cursor position
        menu.exec(self.controls_table.viewport().mapToGlobal(position))

    def show_remap_dialog(self, binding):
        """Show the remapping dialog for a binding"""
        if not self.current_profile:
            return

        from gui.remap_dialog import RemapDialog

        # Use non-modal dialog to avoid Windows keyboard listener suppression
        dialog = RemapDialog(binding.input_code, self.current_profile, self)
        dialog.bindings_changed.connect(self.on_bindings_changed_from_table)
        dialog.setWindowModality(Qt.WindowModality.NonModal)
        dialog.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
        dialog.show()

    def on_binding_remapped(self, binding, new_input_code: str, new_label: str, new_activation_mode: str):
        """Handle binding changes from the remapping dialog"""
        # Update binding's input code
        binding.input_code = new_input_code

        # Update activation mode
        binding.activation_mode = new_activation_mode if new_activation_mode else None

        # Update custom label
        if new_label:
            try:
                from utils.label_overrides import get_override_manager
                override_manager = get_override_manager()

                # Get default label
                default_label = LabelGenerator.generate_action_label(binding.action_name)

                if new_label == default_label:
                    # Remove custom override
                    override_manager.remove_custom_override(binding.action_name)
                    binding.custom_label = None
                else:
                    # Save custom override
                    override_manager.set_custom_override(binding.action_name, new_label)
                    binding.custom_label = new_label

                override_manager.reload()

            except Exception as e:
                logger.error(f"Error updating label: {e}", exc_info=True)

        # Mark profile as modified
        if self.current_profile:
            self.current_profile.mark_modified()

        # Refresh UI
        self.on_profile_modified()

        logger.info(f"Remapped binding: {binding.action_name} to {new_input_code}")

    def on_edit_input_clicked(self, row):
        """Handle Edit button click from table"""
        if not self.current_profile or row >= len(self.all_bindings):
            return

        # Get binding data from the row
        action_map_name, binding = self.all_bindings[row]

        # Show RemapDialog with the current binding
        from gui.remap_dialog import RemapDialog

        # Open in single action mode (hides "Add New Action" section)
        # Pass the binding object directly to handle unmapped actions correctly
        # Use non-modal dialog to avoid Windows keyboard listener suppression
        dialog = RemapDialog(binding.input_code, self.current_profile, self, single_action_mode=True, binding=binding)
        dialog.bindings_changed.connect(self.on_bindings_changed_from_table)
        dialog.setWindowModality(Qt.WindowModality.NonModal)
        dialog.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
        dialog.show()

    def on_bindings_changed_from_table(self, bindings: list):
        """Handle binding changes from RemapDialog opened via table Edit button"""
        # Bindings are already modified by RemapDialog, just refresh UI
        logger.info(f"on_bindings_changed_from_table: Received {len(bindings)} bindings: {[(b.action_name, b.input_code) for b in bindings]}")

        # Mark profile as modified
        if self.current_profile:
            # First pass: remove any duplicate bindings (by object identity) from each action map
            # This can happen if a binding is somehow added multiple times
            for action_map in self.current_profile.action_maps:
                seen_bindings = {}  # Map from binding id to binding object
                bindings_to_keep = []

                for binding in action_map.actions:
                    binding_id = id(binding)
                    if binding_id not in seen_bindings:
                        seen_bindings[binding_id] = binding
                        bindings_to_keep.append(binding)
                    else:
                        logger.warning(f"on_bindings_changed_from_table: Removing duplicate binding {binding.action_name} from {action_map.name}")

                # If we removed any duplicates, update the action map
                if len(bindings_to_keep) < len(action_map.actions):
                    logger.warning(f"on_bindings_changed_from_table: Found and removed {len(action_map.actions) - len(bindings_to_keep)} duplicates in {action_map.name}")
                    action_map.actions = bindings_to_keep

            # Second pass: check if any bindings are not in the profile and add them
            # (This handles multi-binding scenarios)
            for binding in bindings:
                binding_found = False
                binding_location = None

                for action_map in self.current_profile.action_maps:
                    for i, existing_binding in enumerate(action_map.actions):
                        if existing_binding is binding:
                            binding_found = True
                            binding_location = (action_map.name, i)
                            logger.debug(f"on_bindings_changed_from_table: Found binding at {action_map.name}[{i}]")
                            break
                    if binding_found:
                        break

                # If binding is not in profile, add it to the appropriate action map
                if not binding_found:
                    logger.info(f"on_bindings_changed_from_table: Binding {binding.action_name} not found in profile, attempting to add...")
                    # Find the correct action map for this action
                    added = False
                    for action_map in self.current_profile.action_maps:
                        for existing_binding in action_map.actions:
                            if existing_binding.action_name == binding.action_name:
                                # Add the new binding to this action map
                                action_map.actions.append(binding)
                                logger.info(f"on_bindings_changed_from_table: Added binding {binding.action_name} to {action_map.name}")
                                added = True
                                break
                        if added:
                            break

                    if not added:
                        logger.warning(f"on_bindings_changed_from_table: Could not find actionmap for action {binding.action_name}")

            self.current_profile.mark_modified()
            logger.info(f"on_bindings_changed_from_table: Marked profile as modified")

        # Reload override manager cache
        try:
            from utils.label_overrides import get_override_manager
            override_manager = get_override_manager()
            override_manager.reload()
        except Exception as e:
            logger.error(f"Error reloading override manager: {e}", exc_info=True)

        # Refresh UI
        self.on_profile_modified()

    def edit_label_from_context(self, binding):
        """Edit label from context menu (simpler than full remap dialog)"""
        from PyQt6.QtWidgets import QInputDialog

        current_label = LabelGenerator.get_action_label(binding.action_name, binding)

        new_label, ok = QInputDialog.getText(
            self,
            "Edit Control Label",
            f"Edit label for {binding.action_name}:",
            text=current_label
        )

        if ok and new_label != current_label:
            try:
                from utils.label_overrides import get_override_manager
                override_manager = get_override_manager()

                # Get default label
                default_label = LabelGenerator.generate_action_label(binding.action_name)

                if new_label == default_label or not new_label:
                    # Remove custom override
                    override_manager.remove_custom_override(binding.action_name)
                    binding.custom_label = None
                else:
                    # Save custom override
                    override_manager.set_custom_override(binding.action_name, new_label)
                    binding.custom_label = new_label

                override_manager.reload()

                # Mark profile as modified (for label changes)
                if self.current_profile:
                    self.current_profile.mark_modified()

                # Refresh UI
                self.on_profile_modified()

                logger.info(f"Updated label for {binding.action_name}: {new_label}")

            except Exception as e:
                logger.error(f"Error updating label: {e}", exc_info=True)
                QMessageBox.warning(self, "Error", f"Failed to update label: {str(e)}")

    def on_item_double_clicked(self, item):
        """Handle double-click - prepare for editing"""
        if item.column() == 2:  # Label column
            # Get the binding data to store the default label for comparison
            binding_data = item.data(Qt.ItemDataRole.UserRole)
            if binding_data:
                action_map_name, binding = binding_data
                # Store the DEFAULT label (without ANY custom override - not from file, not from binding)
                # We need to get the true default: global override or auto-generated
                # To do this, we use use_override=False to skip the override manager
                self._editing_default_text = LabelGenerator.generate_action_label(binding.action_name)

                # Check if there's a global override (not custom)
                try:
                    from utils.label_overrides import get_override_manager
                except ImportError:
                    from ..utils.label_overrides import get_override_manager

                override_manager = get_override_manager()
                global_overrides = override_manager.load_global_overrides()
                if binding.action_name in global_overrides:
                    self._editing_default_text = global_overrides[binding.action_name]

                logger.debug(f"on_item_double_clicked: default='{self._editing_default_text}', current='{item.text()}'")
            else:
                self._editing_default_text = None

    def on_cell_edited(self, item):
        """Handle cell editing (only column 2: Label is editable)"""
        if item.column() != 2:  # Only Label column is editable
            return

        # Get the new text from the editor
        new_label = item.text().strip()

        # Get stored default text
        default_text = self._editing_default_text

        # Clear editing state
        self._editing_item = None
        self._editing_default_text = None

        # Get the binding data stored in the item
        binding_data = item.data(Qt.ItemDataRole.UserRole)
        if not binding_data:
            self.statusBar().showMessage("Error: Could not save label override")
            return

        action_map_name, binding = binding_data

        # Debug logging
        logger.debug(f"on_cell_edited: action={binding.action_name}, new_label='{new_label}', default='{default_text}'")

        # Check if user deleted everything OR if the new text equals the default
        # In both cases, remove the custom override
        if not new_label or new_label == default_text:
            # Clear the custom_label field to fall back to default
            binding.custom_label = None

            # Get the fallback label
            fallback_label = LabelGenerator.get_action_label(binding.action_name, binding)
            logger.debug(f"Label cleared or matches default, falling back to: '{fallback_label}'")

            # Update the table to show the fallback label
            self.controls_table.blockSignals(True)
            item.setText(fallback_label)
            self.controls_table.blockSignals(False)

            # Remove from the override file
            try:
                # Import with error handling for different execution contexts
                try:
                    from utils.label_overrides import get_override_manager
                except ImportError:
                    from ..utils.label_overrides import get_override_manager

                override_manager = get_override_manager()
                override_manager.remove_custom_override(binding.action_name)

                if not new_label:
                    self.statusBar().showMessage(f"Custom label removed for '{binding.action_name}' - using default: '{fallback_label}'")
                else:
                    self.statusBar().showMessage(f"Label matches default for '{binding.action_name}' - no custom override needed")
                logger.info(f"Removed custom override for '{binding.action_name}', using default: '{fallback_label}'")
            except Exception as e:
                self.statusBar().showMessage(f"Warning: Label restored but couldn't update override file: {str(e)}")
                logger.error(f"Error removing label override: {e}", exc_info=True)

            # Update Device View widget
            if self.pdf_device_widget:
                self.pdf_device_widget.load_profile(self.current_profile)

            # Force reload override manager cache and repopulate the table
            override_manager.reload()
            self.populate_controls_table()
            self.apply_filters()
            return

        # New label is different from default - save as custom override
        try:
            # Import with error handling for different execution contexts
            try:
                from utils.label_overrides import get_override_manager
            except ImportError:
                from ..utils.label_overrides import get_override_manager

            override_manager = get_override_manager()
            override_manager.set_custom_override(binding.action_name, new_label)

            # Update binding's custom_label field
            binding.custom_label = new_label

            # Update the table item to show the new custom label
            self.controls_table.blockSignals(True)
            item.setText(new_label)
            self.controls_table.blockSignals(False)

            # Update Device View widget
            if self.pdf_device_widget:
                self.pdf_device_widget.load_profile(self.current_profile)

            self.statusBar().showMessage(f"Custom label saved: '{binding.action_name}' → '{new_label}'")
            logger.info(f"Saved custom override for '{binding.action_name}': '{new_label}'")

            # Force reload override manager cache and repopulate the table
            override_manager.reload()
            self.populate_controls_table()
            self.apply_filters()
        except Exception as e:
            self.statusBar().showMessage(f"Error saving label override: {str(e)}")
            logger.error(f"Error saving label override: {e}", exc_info=True)

    def export_csv(self):
        """Export profile to CSV"""
        if not self.current_profile:
            return

        # Determine default directory with preference order:
        # 1. Star Citizen profiles directory (if configured and exists)
        # 2. Source profile directory (if available)
        # 3. Current directory
        default_dir = ""

        sc_profiles_dir = self.settings.get_sc_profiles_directory()
        if sc_profiles_dir and os.path.exists(sc_profiles_dir):
            default_dir = sc_profiles_dir
        elif self.current_profile.source_xml_path:
            default_dir = os.path.dirname(self.current_profile.source_xml_path)

        default_path = os.path.join(default_dir, f"{self.current_profile.profile_name}.csv") if default_dir else f"{self.current_profile.profile_name}.csv"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to CSV",
            default_path,
            "CSV Files (*.csv)"
        )

        if not file_path:
            return

        try:
            import csv
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)

                # Determine columns based on detailed view setting
                is_detailed = self.show_detailed_checkbox.isChecked()

                if is_detailed:
                    # Write header for detailed view (all columns)
                    writer.writerow(["Action Category", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"])
                    visible_cols = [0, 1, 2, 3, 4, 5]
                else:
                    # Write header for default view (4 columns)
                    writer.writerow(["Action Category", "Action", "Input Label", "Device"])
                    visible_cols = [0, 2, 4, 5]  # Action Category, Action (Override), Input Label, Device

                # Write visible rows only
                for row in range(self.controls_table.rowCount()):
                    if not self.controls_table.isRowHidden(row):
                        row_data = []
                        for col in visible_cols:
                            item = self.controls_table.item(row, col)
                            row_data.append(item.text() if item else "")
                        writer.writerow(row_data)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to CSV: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export CSV:\n{str(e)}")
            self.statusBar().showMessage("CSV export failed")

    def export_pdf(self):
        """Export profile to PDF"""
        if not self.current_profile:
            return

        # Determine default directory with preference order:
        # 1. Star Citizen profiles directory (if configured and exists)
        # 2. Source profile directory (if available)
        # 3. Current directory
        default_dir = ""

        sc_profiles_dir = self.settings.get_sc_profiles_directory()
        if sc_profiles_dir and os.path.exists(sc_profiles_dir):
            default_dir = sc_profiles_dir
        elif self.current_profile.source_xml_path:
            default_dir = os.path.dirname(self.current_profile.source_xml_path)

        default_path = os.path.join(default_dir, f"{self.current_profile.profile_name}.pdf") if default_dir else f"{self.current_profile.profile_name}.pdf"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to PDF",
            default_path,
            "PDF Files (*.pdf)"
        )

        if not file_path:
            return

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(file_path, pagesize=landscape(letter))
            elements = []
            styles = getSampleStyleSheet()

            # Title
            title = Paragraph(f"<b>{self.current_profile.profile_name}</b>", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 0.2 * inch))

            # Profile info
            info_text = f"Devices: {len(self.current_profile.devices)} | "
            info_text += f"Bindings: {len(self.current_profile.get_all_bindings())}"
            info = Paragraph(info_text, styles['Normal'])
            elements.append(info)
            elements.append(Spacer(1, 0.3 * inch))

            # Determine columns based on detailed view setting
            is_detailed = self.show_detailed_checkbox.isChecked()

            if is_detailed:
                # Detailed view header and columns
                table_data = [["Action Category", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"]]
                visible_cols = [0, 1, 2, 3, 4, 5]
                col_widths = [1.3*inch, 1.5*inch, 1.5*inch, 1*inch, 1.5*inch, 1.2*inch]
            else:
                # Default view header and columns
                table_data = [["Action Category", "Action", "Input Label", "Device"]]
                visible_cols = [0, 2, 4, 5]
                col_widths = [2*inch, 3*inch, 2.5*inch, 1.5*inch]

            for row in range(self.controls_table.rowCount()):
                if not self.controls_table.isRowHidden(row):
                    row_data = []
                    for col in visible_cols:
                        item = self.controls_table.item(row, col)
                        row_data.append(item.text() if item else "")
                    table_data.append(row_data)

            # Create table
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))

            elements.append(table)

            # Build PDF
            doc.build(elements)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to PDF: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export PDF:\n{str(e)}")
            self.statusBar().showMessage("PDF export failed")

    def export_word(self):
        """Export profile to Word document"""
        if not self.current_profile:
            return

        # Determine default directory with preference order:
        # 1. Star Citizen profiles directory (if configured and exists)
        # 2. Source profile directory (if available)
        # 3. Current directory
        default_dir = ""

        sc_profiles_dir = self.settings.get_sc_profiles_directory()
        if sc_profiles_dir and os.path.exists(sc_profiles_dir):
            default_dir = sc_profiles_dir
        elif self.current_profile.source_xml_path:
            default_dir = os.path.dirname(self.current_profile.source_xml_path)

        default_path = os.path.join(default_dir, f"{self.current_profile.profile_name}.docx") if default_dir else f"{self.current_profile.profile_name}.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to Word",
            default_path,
            "Word Documents (*.docx)"
        )

        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(self.current_profile.profile_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Profile info
            info = doc.add_paragraph()
            info.add_run(f"Devices: {len(self.current_profile.devices)} | ")
            info.add_run(f"Total Bindings: {len(self.current_profile.get_all_bindings())}")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # Device list
            doc.add_heading("Devices", 2)
            for device in self.current_profile.devices:
                device_name = device.product_name if device.product_name else device.device_type.capitalize()
                doc.add_paragraph(
                    f"{device.device_type.capitalize()} {device.instance}: {device_name}",
                    style='List Bullet'
                )

            doc.add_paragraph()  # Spacer

            # Control bindings table
            doc.add_heading("Control Bindings", 2)

            # Determine columns based on detailed view setting
            is_detailed = self.show_detailed_checkbox.isChecked()

            if is_detailed:
                headers = ["Action Category", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"]
                visible_cols = [0, 1, 2, 3, 4, 5]
            else:
                headers = ["Action Category", "Action", "Input Label", "Device"]
                visible_cols = [0, 2, 4, 5]

            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add data rows
            for row in range(self.controls_table.rowCount()):
                if not self.controls_table.isRowHidden(row):
                    row_cells = table.add_row().cells
                    for i, col in enumerate(visible_cols):
                        item = self.controls_table.item(row, col)
                        row_cells[i].text = item.text() if item else ""

            # Save document
            doc.save(file_path)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to Word: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export Word document:\n{str(e)}")
            self.statusBar().showMessage("Word export failed")

    def export_graphic(self):
        """Export device graphic - delegate to Device View widget"""
        # Check if Device View tab is active
        current_tab = self.tab_widget.currentIndex()
        if current_tab == 1 and self.pdf_device_widget:
            # Device View tab
            self.pdf_device_widget.export_graphic()
        else:
            # Not on Device View tab
            QMessageBox.warning(self, "Export Graphic", "Please switch to the Device View tab first.")

    def on_profile_modified(self):
        """Handle profile modification - update UI to reflect changes"""
        if not self.current_profile:
            return

        # Update window title with modified indicator
        base_title = f"Star Citizen Profile Editor v{self.version}"
        if self.current_profile.is_modified:
            self.setWindowTitle(f"{base_title} - {self.current_profile.profile_name} *")
            # Enable save button when modified
            self.save_profile_btn.setEnabled(True)
        else:
            self.setWindowTitle(f"{base_title} - {self.current_profile.profile_name}")
            # Disable save button when not modified
            self.save_profile_btn.setEnabled(False)

        # Refresh the controls table
        self.populate_controls_table()
        self.apply_filters()

        # Refresh device view if active
        if self.pdf_device_widget and self.pdf_device_widget.current_device:
            self.pdf_device_widget.load_device_pdf()

    def _add_missing_joystick_devices(self):
        """
        Ensure currently connected joystick devices are in the profile.
        Adds any connected joysticks that aren't already in the profile's device list.
        Also applies device-to-joystick mapping configuration from settings if available.
        Fixes Issue #17: Device detection for new/preset profiles.
        """
        if not self.current_profile:
            logger.warning("_add_missing_joystick_devices: No current profile")
            return

        try:
            from src.utils.input_detector import InputDetector
            from src.models.profile_model import Device

            connected_devices = InputDetector.get_available_devices()
            logger.info(f"_add_missing_joystick_devices: Found {len(connected_devices)} total connected devices")
            for i, dev in enumerate(connected_devices):
                logger.debug(f"  Device {i}: type={dev.get('type')}, instance={dev.get('instance')}, name={dev.get('name')}")

            devices_added = False

            # Load device mapping from settings (user's configured device-to-js mappings)
            device_mapping = self.settings.get_device_config() if self.settings else {}
            logger.info(f"_add_missing_joystick_devices: Loaded device mapping: {device_mapping}")

            # Log current profile devices
            logger.debug(f"_add_missing_joystick_devices: Profile currently has {len(self.current_profile.devices)} devices:")
            for i, d in enumerate(self.current_profile.devices):
                logger.debug(f"  Device {i}: type={d.device_type}, instance={d.instance}, name={d.product_name}")

            # Add joystick devices to the profile if not already present
            joystick_devices = [d for d in connected_devices if d.get('type') == 'joystick']
            logger.info(f"_add_missing_joystick_devices: Found {len(joystick_devices)} joystick device(s) to process")

            for device in joystick_devices:
                device_name = device.get('name', f"Joystick {device.get('instance', 1)}")
                logger.debug(f"_add_missing_joystick_devices: Processing joystick: {device_name}")

                # Determine the instance number based on device mapping
                # First, check if this device has a mapped instance (js1, js2, etc.)
                mapped_instance = None
                for js_label, mapped_device_name in device_mapping.items():
                    if mapped_device_name == device_name:
                        # Extract instance number from label like "js1" -> 1
                        try:
                            mapped_instance = int(js_label.replace('js', ''))
                            logger.debug(f"_add_missing_joystick_devices: Device '{device_name}' is mapped to {js_label} (instance {mapped_instance})")
                            break
                        except ValueError:
                            logger.warning(f"_add_missing_joystick_devices: Could not parse instance from {js_label}")

                # Use mapped instance if found, otherwise use device's instance
                instance = mapped_instance if mapped_instance else device.get('instance', 1)
                logger.debug(f"_add_missing_joystick_devices: Using instance {instance} for {device_name}")

                # Check if this joystick is already in the profile
                already_in_profile = any(
                    d.device_type == 'joystick' and d.instance == instance
                    for d in self.current_profile.devices
                )

                if already_in_profile:
                    logger.debug(f"_add_missing_joystick_devices: Device {device_name} (js{instance}) already in profile")
                else:
                    # Add this joystick to the profile
                    joystick_device = Device(
                        device_type='joystick',
                        instance=instance,
                        product_name=device_name
                    )
                    self.current_profile.devices.append(joystick_device)
                    devices_added = True
                    logger.info(f"_add_missing_joystick_devices: Added joystick device: {device_name} (js{instance})")

            if devices_added:
                logger.info(f"_add_missing_joystick_devices: Profile now has {len(self.current_profile.devices)} devices after adding connected joysticks")
            else:
                logger.info(f"_add_missing_joystick_devices: No new joystick devices were added")

        except Exception as e:
            logger.error(f"_add_missing_joystick_devices: Could not add joystick devices to profile: {e}", exc_info=True)
            # Continue anyway - this is not critical

    def _check_and_warn_disconnected_devices(self, profile):
        """Check if profile has disconnected devices and show warning if needed"""
        try:
            # Get connected devices from config tab
            if not hasattr(self, 'config_tab') or not self.config_tab:
                return

            connected_devices = self.config_tab.current_devices
            if not connected_devices:
                return

            # Find joystick devices in profile
            profile_joysticks = [d for d in profile.devices if d.device_type == 'joystick']
            if not profile_joysticks:
                return

            # Check which profile devices are disconnected
            disconnected = []
            for profile_device in profile_joysticks:
                is_connected = False
                device_name = profile_device.product_name if profile_device.product_name else f"Joystick {profile_device.instance}"

                for connected in connected_devices:
                    if connected.get('type') == 'joystick':
                        if connected.get('instance') == profile_device.instance:
                            # Check product name match
                            if profile_device.product_name:
                                connected_name = connected.get('name', '').lower()
                                device_product = profile_device.product_name.lower()
                                if device_product in connected_name or any(
                                    word in connected_name for word in device_product.split()
                                ):
                                    is_connected = True
                                    break
                            else:
                                is_connected = True
                                break

                if not is_connected:
                    disconnected.append(device_name)

            # Show warning if any devices are disconnected
            if disconnected:
                warning_text = "⚠️ Profile Devices Not Connected\n\n"
                warning_text += "The following devices in this profile are not currently connected:\n\n"
                for device_name in disconnected:
                    warning_text += f"  • {device_name}\n"
                warning_text += "\n"
                warning_text += "Impact:\n"
                warning_text += "  • Input detection will NOT work for disconnected devices\n"
                warning_text += "  • You can still manually select inputs from these devices\n"
                warning_text += "  • Use the manual input selection dropdown in the editing interface\n"
                warning_text += "\n"
                warning_text += "You can still view and edit templates for these devices in the Device View tab."

                QMessageBox.information(self, "Disconnected Devices", warning_text)
                logger.info(f"Profile loaded with {len(disconnected)} disconnected device(s): {', '.join(disconnected)}")

        except Exception as e:
            logger.error(f"Error checking disconnected devices: {e}", exc_info=True)

    def on_update_profile_name(self):
        """Handle profile name change from input field"""
        if not self.current_profile:
            return

        new_name = self.profile_name_input.text().strip()

        # Validate name is not empty
        if not new_name:
            QMessageBox.warning(self, "Invalid Name", "Profile name cannot be empty.")
            self.profile_name_input.setText(self.current_profile.profile_name)
            return

        # Check if name actually changed
        if new_name == self.current_profile.profile_name:
            return

        # Update profile name
        self.current_profile.profile_name = new_name

        # Mark as modified
        self.current_profile.mark_modified()

        # Update UI
        self.profile_info_label.setText(f"Profile: {new_name}")
        self.on_profile_modified()

        logger.info(f"Profile name updated to: {new_name}")

    def save_profile(self):
        """Save the modified profile to XML"""
        if not self.current_profile or not self.current_profile.is_modified:
            return

        # Prompt user for profile name
        profile_name, ok = QInputDialog.getText(
            self,
            "Profile Name",
            "Enter a name for your profile:",
            QLineEdit.EchoMode.Normal,
            self.current_profile.profile_name  # Default to current name
        )

        if not ok or not profile_name.strip():
            # User cancelled or entered empty name
            return

        # Update the profile name
        profile_name = profile_name.strip()
        self.current_profile.profile_name = profile_name

        # Generate filename using layout_<name>.xml pattern
        suggested_name = f"layout_{profile_name}.xml"

        # Determine default directory with preference order:
        # 1. Star Citizen profiles directory (if configured and exists)
        # 2. Source profile directory (if available)
        # 3. Current directory
        suggested_dir = ""

        sc_profiles_dir = self.settings.get_sc_profiles_directory()
        if sc_profiles_dir and os.path.exists(sc_profiles_dir):
            suggested_dir = sc_profiles_dir
        elif self.current_profile.source_xml_path:
            suggested_dir = os.path.dirname(self.current_profile.source_xml_path)

        default_path = os.path.join(suggested_dir, suggested_name) if suggested_dir else suggested_name

        # Show save dialog with generated filename
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Profile",
            default_path,
            "XML Files (*.xml);;All Files (*)"
        )

        if not output_path:
            return

        # Export profile to XML
        try:
            from exporters.xml_exporter import export_profile

            # Use original source path if available, otherwise create from scratch
            source_path = self.current_profile.source_xml_path

            if source_path and os.path.exists(source_path):
                success = export_profile(self.current_profile, source_path, output_path)
            else:
                # Create new profile from scratch
                from exporters.xml_exporter import XMLExporter
                success = XMLExporter.create_new_profile(self.current_profile, output_path)

            if success:
                # Mark as saved
                self.current_profile.mark_saved()
                self.current_profile.source_xml_path = output_path

                # Update UI
                self.on_profile_modified()

                # Show success message
                QMessageBox.information(
                    self,
                    "Success",
                    f"Profile saved successfully!\n\nSaved to:\n{output_path}"
                )

                self.statusBar().showMessage(f"Profile saved: {os.path.basename(output_path)}")
                logger.info(f"Saved profile to: {output_path}")

            else:
                QMessageBox.critical(
                    self,
                    "Save Error",
                    "Failed to save profile. Check the log for details."
                )

        except Exception as e:
            logger.error(f"Error saving profile: {e}", exc_info=True)
            QMessageBox.critical(
                self,
                "Save Error",
                f"Failed to save profile:\n{str(e)}"
            )

    def show_help(self):
        """Show the user guide in a dialog with TOC sidebar"""
        logger.info("=== show_help() called ===")
        try:
            # Create help dialog
            logger.info("Creating help dialog...")
            help_dialog = QDialog(self)
            help_dialog.setWindowTitle(f"User Guide - Star Citizen Profile Editor v{self.version}")
            help_dialog.setGeometry(100, 100, 1200, 800)
            logger.info("Help dialog created successfully")

            # Main layout
            logger.info("Creating main layout...")
            main_layout = QVBoxLayout()
            help_dialog.setLayout(main_layout)
            logger.info("Main layout set")

            # Content area layout (horizontal: sidebar + browser)
            logger.info("Creating content layout...")
            content_layout = QHBoxLayout()

            # Text browser for displaying markdown as HTML
            logger.info("Creating text browser...")
            browser = QTextBrowser()
            browser.setOpenExternalLinks(True)
            logger.info("Text browser created")

            # Create sidebar for table of contents
            logger.info("Creating TOC widget...")
            toc_widget = QListWidget()
            toc_widget.setMaximumWidth(250)
            logger.info("Connecting TOC signal...")
            toc_widget.itemClicked.connect(lambda item: self.navigate_to_section(browser, item.data(Qt.ItemDataRole.UserRole)))
            logger.info("TOC signal connected")
            content_layout.addWidget(toc_widget)
            content_layout.addWidget(browser)
            logger.info("Widgets added to content layout")

            main_layout.addLayout(content_layout)
            logger.info("Content layout added to main layout")

            # Load the README.md file
            logger.info("Finding README.md file...")
            user_guide_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                "README.md"
            )
            logger.info(f"README.md path: {user_guide_path}")

            try:
                logger.info("Opening README.md...")
                with open(user_guide_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                logger.info(f"README.md loaded: {len(markdown_content)} characters")

                # Extract headings for TOC and convert markdown to HTML
                logger.info("Extracting headings...")
                headings = self.extract_headings(markdown_content)
                logger.info(f"Found {len(headings)} headings")

                # Populate TOC
                logger.info("Populating TOC...")
                for heading_level, heading_text in headings:
                    logger.debug(f"Adding TOC item: level={heading_level}, text='{heading_text}'")
                    item = QListWidgetItem(heading_text)
                    anchor_id = self.create_anchor_id(heading_text)
                    logger.debug(f"Anchor ID: {anchor_id}")
                    item.setData(Qt.ItemDataRole.UserRole, anchor_id)
                    # Indent sub-headings
                    if heading_level > 2:
                        item.setText("  " * (heading_level - 2) + heading_text)
                    toc_widget.addItem(item)
                logger.info(f"TOC populated with {toc_widget.count()} items")

                # Convert markdown to HTML
                logger.info("Converting markdown to HTML...")
                html_content = self.markdown_to_html(markdown_content)
                logger.info(f"HTML generated: {len(html_content)} characters")
                logger.info("Setting HTML content in browser...")
                browser.setHtml(html_content)
                logger.info("HTML content set successfully")

            except FileNotFoundError as e:
                logger.error(f"README.md not found: {e}")
                browser.setHtml("<h1>User Guide Not Found</h1><p>The README.md file could not be found.</p>")
            except Exception as e:
                logger.error(f"Error loading guide: {e}", exc_info=True)
                browser.setHtml(f"<h1>Error Loading Guide</h1><p>{str(e)}</p>")

            # Close button
            logger.info("Creating close button...")
            button_layout = QHBoxLayout()
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(help_dialog.close)
            button_layout.addStretch()
            button_layout.addWidget(close_btn)
            main_layout.addLayout(button_layout)
            logger.info("Close button added")

            logger.info("Showing help dialog...")
            help_dialog.exec()
            logger.info("Help dialog closed")

        except Exception as e:
            logger.error(f"Fatal error in show_help(): {e}", exc_info=True)
            QMessageBox.critical(self, "Help Dialog Error", f"Failed to show help dialog:\n{str(e)}")

    def markdown_to_html(self, markdown_text):
        """Convert markdown to HTML (basic implementation)"""
        # Get theme-aware colors from the application palette
        palette = self.palette()
        text_color = palette.color(palette.ColorRole.Text).name()
        base_color = palette.color(palette.ColorRole.Base).name()
        link_color = palette.color(palette.ColorRole.Link).name()

        # Increase font sizes for better readability
        html = "<html><head><style>"
        html += f"body {{ font-family: Segoe UI, Arial, sans-serif; line-height: 1.8; padding: 20px; font-size: 16px; color: {text_color}; background-color: {base_color}; }}"
        html += f"h1 {{ color: {link_color}; border-bottom: 3px solid {link_color}; padding-bottom: 10px; font-size: 36px; font-weight: bold; margin-top: 20px; }}"
        html += f"h2 {{ color: {link_color}; border-bottom: 2px solid {link_color}; padding-bottom: 5px; margin-top: 30px; font-size: 28px; font-weight: bold; }}"
        html += f"h3 {{ color: {link_color}; margin-top: 20px; font-size: 22px; font-weight: bold; }}"
        html += f"h4 {{ color: {link_color}; margin-top: 15px; font-size: 18px; font-weight: bold; }}"
        html += f"p {{ font-size: 16px; margin: 10px 0; color: {text_color}; }}"
        html += f"li {{ font-size: 16px; margin: 5px 0; color: {text_color}; }}"
        html += f"a {{ color: {link_color}; text-decoration: underline; font-weight: 500; }}"
        html += f"a:hover {{ text-decoration: underline; opacity: 0.8; cursor: pointer; }}"
        html += f"code {{ background-color: rgba(0,0,0,0.05); padding: 2px 6px; border-radius: 3px; font-family: 'Courier New', monospace; font-size: 15px; color: {text_color}; }}"
        html += f"pre {{ background-color: rgba(0,0,0,0.05); padding: 10px; border-radius: 5px; overflow-x: auto; font-size: 15px; color: {text_color}; }}"
        html += f"ul {{ margin-left: 20px; font-size: 16px; }}"
        html += f"ol {{ margin-left: 20px; font-size: 16px; }}"
        html += f"strong {{ color: {text_color}; font-weight: bold; }}"
        html += f"blockquote {{ border-left: 4px solid {link_color}; padding-left: 15px; color: {text_color}; font-style: italic; font-size: 16px; }}"
        html += f"table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}"
        html += f"th, td {{ border: 1px solid {link_color}; padding: 10px; text-align: left; font-size: 16px; color: {text_color}; }}"
        html += f"th {{ background-color: {link_color}; color: white; font-weight: bold; }}"
        html += f"tr:nth-child(even) {{ background-color: rgba(0,0,0,0.03); }}"
        html += "</style></head><body>"

        lines = markdown_text.split('\n')
        in_code_block = False
        in_list = False
        in_blockquote = False

        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    html += "</pre>"
                    in_code_block = False
                else:
                    html += "<pre><code>"
                    in_code_block = True
                continue

            if in_code_block:
                html += line + "\n"
                continue

            # Skip image badges (![...])
            if line.strip().startswith('!['):
                continue

            # Headers (with anchor IDs for links)
            if line.startswith('# '):
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                header_text = line[2:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h1 id='{anchor_id}'>{header_text}</h1>"
            elif line.startswith('## '):
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                header_text = line[3:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h2 id='{anchor_id}'>{header_text}</h2>"
            elif line.startswith('### '):
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                header_text = line[4:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h3 id='{anchor_id}'>{header_text}</h3>"
            elif line.startswith('#### '):
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                header_text = line[5:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h4 id='{anchor_id}'>{header_text}</h4>"
            # Blockquotes
            elif line.strip().startswith('> '):
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if not in_blockquote:
                    html += "<blockquote>"
                    in_blockquote = True
                blockquote_text = line.strip()[2:].strip()
                html += f"{self.format_inline_markdown(blockquote_text)}<br/>"
            # Ordered lists
            elif line.strip() and line.strip()[0].isdigit() and '. ' in line.strip()[:4]:
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                if in_list != 'ol':
                    if in_list:
                        html += f"</{in_list}>"
                    html += "<ol>"
                    in_list = 'ol'
                # Extract the list item text (after the number and period)
                item_text = line.strip().split('. ', 1)[1] if '. ' in line.strip() else line.strip()
                html += f"<li>{self.format_inline_markdown(item_text)}</li>"
            # Unordered lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                if in_list != 'ul':
                    if in_list:
                        html += f"</{in_list}>"
                    html += "<ul>"
                    in_list = 'ul'
                html += f"<li>{self.format_inline_markdown(line.strip()[2:])}</li>"
            # Horizontal rule
            elif line.strip() == '---':
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                html += "<hr/>"
            # Empty line
            elif not line.strip():
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                # Don't add extra breaks, just skip empty lines
            # Regular paragraph
            else:
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                if in_blockquote:
                    html += "</blockquote>"
                    in_blockquote = False
                html += f"<p>{self.format_inline_markdown(line)}</p>"

        if in_list:
            html += f"</{in_list}>"
        if in_blockquote:
            html += "</blockquote>"

        html += "</body></html>"
        return html

    def create_anchor_id(self, header_text):
        """Create an anchor ID from header text (GitHub-style)"""
        import re
        # Convert to lowercase
        anchor = header_text.lower()
        # Replace spaces with hyphens
        anchor = anchor.replace(' ', '-')
        # Remove special characters except hyphens
        anchor = re.sub(r'[^\w\-]', '', anchor)
        return anchor

    def format_inline_markdown(self, text):
        """Format inline markdown (bold, italic, code, links)"""
        import re

        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Inline code
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        # Links (handle both URLs and anchors)
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)

        return text

    def extract_headings(self, markdown_text):
        """Extract all headings from markdown text with their levels"""
        headings = []
        lines = markdown_text.split('\n')

        for line in lines:
            if line.startswith('# '):
                headings.append((1, line[2:].strip()))
            elif line.startswith('## '):
                headings.append((2, line[3:].strip()))
            elif line.startswith('### '):
                headings.append((3, line[4:].strip()))
            elif line.startswith('#### '):
                headings.append((4, line[5:].strip()))

        return headings

    def navigate_to_section(self, browser, anchor_id):
        """Navigate to a section in the browser by anchor ID"""
        if anchor_id:
            browser.scrollToAnchor(anchor_id)

    def open_discord_link(self, event):
        """Open Discord invite link in browser"""
        discord_url = "https://discord.gg/BNzRegKZ7k"
        QDesktopServices.openUrl(QUrl(discord_url))

    def open_paypal_donation(self, event):
        """Open PayPal donation link in browser"""
        paypal_url = "https://paypal.me/RighteousKill"
        QDesktopServices.openUrl(QUrl(paypal_url))

    def open_venmo_donation(self, event):
        """Open Venmo donation link in browser"""
        venmo_url = "https://venmo.com/u/Amr-Abouelleil"
        QDesktopServices.openUrl(QUrl(venmo_url))


if __name__ == "__main__":
    from PyQt6.QtWidgets import QApplication

    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())